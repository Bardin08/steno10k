from __future__ import annotations

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph

from steno10k.lib.docx import _split_inline, render_markdown_to_docx


def test_plain_text_is_single_plain_run() -> None:
    assert _split_inline("hello world") == [("hello world", "plain")]


def test_bold_span() -> None:
    assert _split_inline("a **b** c") == [("a ", "plain"), ("b", "bold"), (" c", "plain")]


def test_italic_span_with_asterisk() -> None:
    assert _split_inline("a *b* c") == [("a ", "plain"), ("b", "italic"), (" c", "plain")]


def test_italic_span_with_underscore() -> None:
    assert _split_inline("a _b_ c") == [("a ", "plain"), ("b", "italic"), (" c", "plain")]


def test_code_span() -> None:
    assert _split_inline("see `x = 1` now") == [
        ("see ", "plain"),
        ("x = 1", "code"),
        (" now", "plain"),
    ]


def test_bold_takes_precedence_over_italic() -> None:
    assert _split_inline("**bold**") == [("bold", "bold")]


def test_stray_marker_is_literal() -> None:
    assert _split_inline("a * b") == [("a * b", "plain")]


def test_unclosed_bold_is_literal() -> None:
    assert _split_inline("a **b c") == [("a **b c", "plain")]


def test_empty_string() -> None:
    assert _split_inline("") == []


def test_code_content_is_not_reparsed() -> None:
    assert _split_inline("`**not bold**`") == [("**not bold**", "code")]


def _style_name(p: Paragraph) -> str:
    assert p.style is not None
    return str(p.style.name)


def _paras(doc: DocumentObject) -> list[tuple[str, str]]:
    return [(_style_name(p), p.text) for p in doc.paragraphs]


def test_headings_levels() -> None:
    doc = Document()
    render_markdown_to_docx("# H1\n## H2\n### H3\n#### H4", doc)
    styles = [_style_name(p) for p in doc.paragraphs if p.text]
    assert styles == ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]


def test_plain_paragraph() -> None:
    doc = Document()
    render_markdown_to_docx("Just a sentence.", doc)
    assert ("Normal", "Just a sentence.") in _paras(doc)


def test_bullet_list() -> None:
    doc = Document()
    render_markdown_to_docx("- one\n- two", doc)
    assert [p.text for p in doc.paragraphs if _style_name(p) == "List Bullet"] == ["one", "two"]


def test_nested_bullet() -> None:
    doc = Document()
    render_markdown_to_docx("- top\n  - child", doc)
    assert any(_style_name(p) == "List Bullet 2" and p.text == "child" for p in doc.paragraphs)


def test_numbered_list() -> None:
    doc = Document()
    render_markdown_to_docx("1. first\n2. second", doc)
    numbers = [p.text for p in doc.paragraphs if _style_name(p) == "List Number"]
    assert numbers == ["first", "second"]


def test_horizontal_rule_is_skipped() -> None:
    doc = Document()
    render_markdown_to_docx("a\n\n---\n\nb", doc)
    texts = [p.text for p in doc.paragraphs if p.text]
    assert "---" not in texts and "a" in texts and "b" in texts


def test_bold_run_in_paragraph() -> None:
    doc = Document()
    render_markdown_to_docx("a **b** c", doc)
    para = next(p for p in doc.paragraphs if p.text == "a b c")
    assert [r.text for r in para.runs if r.bold] == ["b"]


def test_blank_lines_do_not_create_empty_paragraphs() -> None:
    doc = Document()
    render_markdown_to_docx("a\n\n\nb", doc)
    assert [p.text for p in doc.paragraphs if p.text] == ["a", "b"]
