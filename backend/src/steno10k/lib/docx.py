from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document as new_document
from docx.document import Document
from docx.text.paragraph import Paragraph

from steno10k.contracts.errors import ErrorLog

log = logging.getLogger("steno10k.docx")

_INLINE_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("code", re.compile(r"`([^`]+)`")),
    ("bold", re.compile(r"\*\*([^*]+)\*\*")),
    ("italic", re.compile(r"\*([^*]+)\*")),
    ("italic", re.compile(r"_([^_]+)_")),
]


def _split_inline(text: str) -> list[tuple[str, str]]:
    """Tokenize one line into (text, style) runs: plain|bold|italic|code."""
    if not text:
        return []
    best: tuple[int, int, str, str] | None = None
    for style, pat in _INLINE_PATTERNS:
        m = pat.search(text)
        if m is None:
            continue
        if best is None or m.start() < best[0]:
            best = (m.start(), m.end(), style, m.group(1))
    if best is None:
        return [(text, "plain")]
    start, end, style, inner = best
    runs: list[tuple[str, str]] = []
    if start > 0:
        runs.extend(_split_inline(text[:start]))
    runs.append((inner, style))
    if end < len(text):
        runs.extend(_split_inline(text[end:]))
    return runs


_HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_HR_RE = re.compile(r"^\s*---+\s*$")


def _apply_inline(paragraph: Paragraph, text: str) -> None:
    for run_text, style in _split_inline(text):
        run = paragraph.add_run(run_text)
        if style == "bold":
            run.bold = True
        elif style == "italic":
            run.italic = True
        elif style == "code":
            run.font.name = "Courier New"


def render_markdown_to_docx(markdown: str, doc: Document) -> None:
    """Append a supported Markdown subset to `doc`.

    Supports ATX headings (#-####), `-`/`*` bullets (nested by indent), numbered
    lists, `**bold**`/`*italic*`/`_italic_`/`` `code` `` spans, plain paragraphs.
    `---` rules and blank lines produce nothing; malformed markup degrades to text.
    """
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip() or _HR_RE.match(line):
            continue
        m = _HEADING_RE.match(line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            continue
        m = _BULLET_RE.match(line)
        if m:
            style = "List Bullet 2" if len(m.group(1)) >= 2 else "List Bullet"
            _apply_inline(doc.add_paragraph(style=style), m.group(2).strip())
            continue
        m = _NUMBERED_RE.match(line)
        if m:
            _apply_inline(doc.add_paragraph(style="List Number"), m.group(1).strip())
            continue
        _apply_inline(doc.add_paragraph(), line.strip())


def _build_one(source: Path, dst: Path, force: bool) -> None:
    """Render one markdown source into one .docx at dst. Assumes source exists."""
    if dst.exists() and not force:
        return
    doc = new_document()
    render_markdown_to_docx(source.read_text(encoding="utf-8"), doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def build_bundle(
    pairs: list[tuple[Path, Path]],
    force: bool,
    err_log: ErrorLog,
) -> None:
    """Render each (markdown_path -> docx_path) pair independently.

    A missing source is skipped silently (not an error). A render failure on one
    pair is logged to `err_log` and does not abort the others (per-file isolation).
    """
    for source, dst in pairs:
        if not source.exists():
            log.info("bundle: skipping %s (source %s not found)", dst.name, source.name)
            continue
        try:
            _build_one(source, dst, force)
        except Exception as e:  # per-file isolation boundary (lxml/zip errors, etc.)
            err_log.log("bundle", dst.name, e)
